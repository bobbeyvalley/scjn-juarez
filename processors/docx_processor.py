# processors/docx_processor.py
"""
Procesador para archivos DOCX y DOC
"""

import docx
from pathlib import Path
from typing import Tuple

class DOCXProcessor:
    def extraer_contenido(self, ruta_archivo: Path) -> Tuple[str, str]:
        """
        Extrae contenido de un archivo DOCX/DOC
        
        Returns:
            Tuple[str, str]: (contenido_texto, tipo_contenido)
        """
        try:
            doc = docx.Document(ruta_archivo)
            texto_completo = []
            
            # Extraer párrafos
            for para in doc.paragraphs:
                texto_completo.append(para.text)
            
            # Extraer tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        texto_completo.append(cell.text)
            
            contenido = "\n".join(texto_completo)
            return contenido, "text"
            
        except Exception as e:
            raise Exception(f"Error procesando documento DOCX: {str(e)}")
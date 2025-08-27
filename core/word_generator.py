# core/word_generator.py
"""
Generador de documentos Word profesionales para proyectos de sentencia SCJN
"""

import re
from pathlib import Path
from datetime import datetime
from typing import Optional, Dict, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

class SCJNWordGenerator:
    """Generador de documentos Word con formato profesional para la SCJN"""
    
    def __init__(self):
        self.doc = None
        self.estilos_configurados = False

    #a ver si esto no rompe todo
    def crear_documento_desde_markdown(self, contenido_markdown: str, expediente: str,
                                 archivo_salida: Path, nombre_ministro: str,
                                 nombre_secretario: str, contexto_caso: dict,
                                 nombre_secretario_aux: str = None,
                                 colaboradores: list = None) -> Path:
        """
        Convierte contenido Markdown a documento Word profesional
        
        Args:
            contenido_markdown: Contenido en formato Markdown
            expediente: Número de expediente para el header
            archivo_salida: Ruta donde guardar el documento Word
            
        Returns:
            Path: Ruta del archivo Word generado
        """
            
        self.doc = Document()
        self._configurar_estilos_profesionales()
        #self._crear_estructura_dos_secciones()

        # Primera sección - página de presentación  
        self._agregar_header_scjn(
            expediente=expediente,
            nombre_ministro=nombre_ministro,
            nombre_secretario=nombre_secretario,
            contexto_caso=contexto_caso,
            nombre_secretario_aux=nombre_secretario_aux,
            colaboradores=colaboradores
        )

        # Segunda sección - repetir header + contenido
        self._agregar_header_segunda_seccion(
            expediente=expediente,
            nombre_ministro=nombre_ministro,
            nombre_secretario=nombre_secretario,
            contexto_caso=contexto_caso,
            nombre_secretario_aux=nombre_secretario_aux
        )
        
        # Procesar contenido Markdown
        print("🔍 DEBUG - Iniciando procesamiento de Markdown...")
        self._procesar_markdown(contenido_markdown)
        print("🔍 DEBUG - Markdown procesado") 

        # Agregar footer
        #print("🔍 DEBUG - Agregando footer...")
        self._agregar_footer()
        #print("🔍 DEBUG - Footer agregado")     
    
        # Guardar documento
        print(f"🔍 DEBUG - Guardando documento en: {archivo_salida}")
        self.doc.save(archivo_salida)
        print(f"🔍 DEBUG - Archivo guardado exitosamente: {archivo_salida.exists()}")
        return archivo_salida
        

    def _configurar_estilos_profesionales(self):
        """Configura estilos profesionales para documentos SCJN"""
        if self.estilos_configurados:
            return
            
        styles = self.doc.styles
        
        # Estilo para título principal
        try:
            titulo_principal = styles.add_style('TituloPrincipal', WD_STYLE_TYPE.PARAGRAPH)
        except:
            titulo_principal = styles['TituloPrincipal']
        
        titulo_principal.font.name = 'Arial'
        titulo_principal.font.size = Pt(13)
        titulo_principal.font.bold = True
        titulo_principal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        titulo_principal.paragraph_format.space_after = Pt(13)
        
        # Estilo para subtítulos nivel 2
        try:
            subtitulo_h2 = styles.add_style('SubtituloH2', WD_STYLE_TYPE.PARAGRAPH)
        except:
            subtitulo_h2 = styles['SubtituloH2']
            
        subtitulo_h2.font.name = 'Arial'
        subtitulo_h2.font.size = Pt(13)
        subtitulo_h2.font.bold = True
        subtitulo_h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        subtitulo_h2.paragraph_format.space_before = Pt(12)
        subtitulo_h2.paragraph_format.space_after = Pt(6)
        
        # Estilo para subtítulos nivel 3
        try:
            subtitulo_h3 = styles.add_style('SubtituloH3', WD_STYLE_TYPE.PARAGRAPH)
        except:
            subtitulo_h3 = styles['SubtituloH3']
            
        subtitulo_h3.font.name = 'Arial'
        subtitulo_h3.font.size = Pt(13)
        subtitulo_h3.font.bold = True
        subtitulo_h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        subtitulo_h3.paragraph_format.space_before = Pt(6)
        subtitulo_h3.paragraph_format.space_after = Pt(3)
        
        # Estilo para texto normal
        try:
            texto_normal = styles.add_style('TextoNormal', WD_STYLE_TYPE.PARAGRAPH)
        except:
            texto_normal = styles['TextoNormal']
            
        texto_normal.font.name = 'Arial'
        texto_normal.font.size = Pt(13)
        texto_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        texto_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        texto_normal.paragraph_format.space_after = Pt(6)
        texto_normal.paragraph_format.first_line_indent = Inches(0.5)
        
        # Estilo para citas textuales (indentadas)
        try:
            cita_textual = styles.add_style('CitaTextual', WD_STYLE_TYPE.PARAGRAPH)
        except:
            cita_textual = styles['CitaTextual']
            
        cita_textual.font.name = 'Arial'
        cita_textual.font.size = Pt(13)
        cita_textual.font.italic = True
        cita_textual.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        cita_textual.paragraph_format.left_indent = Inches(0.5)
        cita_textual.paragraph_format.right_indent = Inches(0.5)
        cita_textual.paragraph_format.space_before = Pt(3)
        cita_textual.paragraph_format.space_after = Pt(3)
        
        self.estilos_configurados = True

    def _configurar_pagina(self):
        """Configura márgenes y orientación de página"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.5)
            section.right_margin = Inches(1.0)

    def _crear_estructura_dos_secciones(self):
        """Configura documento con dos secciones separadas"""
        # Sección 1 - sin header/footer (página de presentación)
        section1 = self.doc.sections[0]
        section1.top_margin = Inches(1.0)
        section1.bottom_margin = Inches(1.0)
        section1.left_margin = Inches(1.5)
        section1.right_margin = Inches(1.0)
        
        # Agregar section break después de la primera sección
        self.doc.add_section()
        
        # Sección 2 - con header/footer
        section2 = self.doc.sections[1]
        section2.top_margin = Inches(1.0)
        section2.bottom_margin = Inches(1.0) 
        section2.left_margin = Inches(1.5)
        section2.right_margin = Inches(1.0)
        
        # Configurar header para sección 2
        header = section2.header
        header_para = header.paragraphs[0]
        header_para.text = "[RECURSO - EXPEDIENTE]"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_para.runs[0]
        run.font.name = 'Arial'
        run.font.size = Pt(13)
        run.font.bold = True
        
        # Configurar footer para sección 2 (numeración)
        footer = section2.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Agregar número de página
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        
        fldChar1 = parse_xml(r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        instrText = parse_xml(r'<w:instrText xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> PAGE </w:instrText>')
        fldChar2 = parse_xml(r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        
        footer_para._p.append(fldChar1)
        footer_para._p.append(instrText) 
        footer_para._p.append(fldChar2)
        
        # Configurar numeración para que empiece en página 2
        section2.start_type = 2  # Página continua
        section2.page_start = 2

    def _agregar_header_scjn(self, expediente: str, nombre_ministro: str, 
                        nombre_secretario: str, contexto_caso: dict,
                        nombre_secretario_aux: str = None, 
                        colaboradores: list = None):
        """Agrega header de primera página con formato especial"""
        
        # SECCIÓN 1 - PÁGINA DE PRESENTACIÓN
        # Tipo de caso y datos (con margen especial)
        caso_para = self.doc.add_paragraph()
        caso_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        caso_para.paragraph_format.left_indent = Inches(3.15)  # ~8cm
        
        expediente_principal = contexto_caso.get('metadata_caso', {}).get('expediente_principal', f'AMPARO DIRECTO EN REVISIÓN {expediente}')
        
        # Línea del caso
        run1 = caso_para.add_run(expediente_principal)
        run1.font.name = 'Arial'
        run1.font.size = Pt(13)
        run1.font.bold = True
        
        caso_para.add_run('\n')
        
        # Línea del quejoso
        run2 = caso_para.add_run('QUEJOSA Y RECURRENTE: ')
        run2.font.name = 'Arial'
        run2.font.size = Pt(13)
        run2.font.bold = True
        
        run3 = caso_para.add_run('[INSERTAR DATOS DEL QUEJOSO]')
        run3.font.name = 'Arial'
        run3.font.size = Pt(13)
        run3.font.bold = True
        run3.font.color.rgb = RGBColor(255, 0, 0)  # Rojo
        
        # Espacio
        self.doc.add_paragraph()
        
        # Personal jurisdiccional
        personal = self.doc.add_paragraph()
        personal.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        personal.add_run(f'PONENTE: MINISTRO {nombre_ministro.upper()}\n')
        personal.add_run(f'SECRETARIO: {nombre_secretario.upper()}')
        
        if nombre_secretario_aux:
            personal.add_run(f'\nSECRETARIO AUXILIAR: {nombre_secretario_aux.upper()}')
        
        # Aplicar formato al párrafo de personal
        for run in personal.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(13)
            run.font.bold = True
        
        # Espacio antes de la tabla
        self.doc.add_paragraph()
        
        # Tabla del índice temático (4x2)
        tabla_indice = self.doc.add_table(rows=2, cols=4)
        tabla_indice.style = 'Table Grid'
        
        # Headers de la tabla
        headers = tabla_indice.rows[0].cells
        headers[0].text = ""  # Vacía
        headers[1].text = "Apartado"
        headers[2].text = "Criterio y decisión"
        headers[3].text = "Págs."
        
        # Formatear headers
        for cell in headers:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(13)
                    run.font.bold = True
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _agregar_header_segunda_seccion(self, expediente: str, nombre_ministro: str, 
                               nombre_secretario: str, contexto_caso: dict,
                               nombre_secretario_aux: str = None):
        """Agrega header repetido en segunda sección"""
        
        # Cambiar a segunda sección
        self.doc.add_page_break()
        
        # Repetir información del caso (CON MARGEN IGUAL QUE PRIMERA PÁGINA)
        caso_para2 = self.doc.add_paragraph()
        caso_para2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        caso_para2.paragraph_format.left_indent = Inches(3.15)  # MISMO MARGEN
        caso_para2.paragraph_format.space_after = Pt(12)
        
        expediente_principal = contexto_caso.get('metadata_caso', {}).get('expediente_principal', f'AMPARO DIRECTO EN REVISIÓN {expediente}')
        
        run1 = caso_para2.add_run(expediente_principal)
        run1.font.name = 'Arial'
        run1.font.size = Pt(13)
        run1.font.bold = True
        
        caso_para2.add_run('\n')
        
        run2 = caso_para2.add_run('QUEJOSA Y RECURRENTE: [INSERTAR DATOS DEL QUEJOSO]')
        run2.font.name = 'Arial'
        run2.font.size = Pt(13) 
        run2.font.bold = True
        
        # Espacio
        self.doc.add_paragraph()
        
        # Repetir personal (FORMATO IGUAL QUE PRIMERA PÁGINA)
        personal2 = self.doc.add_paragraph()
        personal2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        personal2.paragraph_format.space_after = Pt(12)
        
        personal2.add_run(f'PONENTE: MINISTRO {nombre_ministro.upper()}\n')
        personal2.add_run(f'SECRETARIO: {nombre_secretario.upper()}')
        
        if nombre_secretario_aux:
            personal2.add_run(f'\nSECRETARIO AUXILIAR: {nombre_secretario_aux.upper()}')
        
        for run in personal2.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(13)
            run.font.bold = True
        
        # Espacio
        self.doc.add_paragraph()
        
        # Placeholder SENTENCIA
        sentencia_para = self.doc.add_paragraph()
        sentencia_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sentencia_para.paragraph_format.space_before = Pt(24)
        sentencia_para.paragraph_format.space_after = Pt(24)
        
        run_sentencia = sentencia_para.add_run('[SENTENCIA]')
        run_sentencia.font.name = 'Arial'
        run_sentencia.font.size = Pt(13)
        run_sentencia.font.bold = True

    def _procesar_markdown(self, contenido: str):
        """Procesa el contenido Markdown línea por línea - Versión Segura"""
        print(f"🔍 DEBUG - Procesando {len(contenido)} caracteres de Markdown")
        lineas = contenido.split('\n')
        print(f"🔍 DEBUG - Dividido en {len(lineas)} líneas")
        
        i = 0
        max_iteraciones = len(lineas) * 2  # Límite de seguridad
        iteraciones = 0
        
        while i < len(lineas) and iteraciones < max_iteraciones:
            iteraciones += 1
            
            if iteraciones % 100 == 0:  # Debug cada 100 iteraciones
                print(f"🔍 DEBUG - Iteración {iteraciones}, línea {i}/{len(lineas)}")
            
            linea = lineas[i].strip()
            
            if not linea:
                # Línea vacía - agregar espacio
                self.doc.add_paragraph()
                i += 1
                continue
            
            # Headers
            if linea.startswith('### '):
                # H3 - Subsecciones
                titulo = linea[4:].strip()
                p = self.doc.add_paragraph()
                p.style = 'SubtituloH3'
                p.add_run(titulo)
                i += 1
                continue
            elif linea.startswith('## '):
                # H2 - Secciones principales
                titulo = linea[3:].strip()
                p = self.doc.add_paragraph()
                p.style = 'SubtituloH2'
                p.add_run(titulo)
                i += 1
                continue
            elif linea.startswith('# '):
                # H1 - Título principal (saltar)
                i += 1
                continue
            
            # Líneas de separación
            elif linea.startswith('---') or linea.startswith('==='):
                self.doc.add_paragraph()
                i += 1
                continue
            
            # Texto con formato especial
            elif linea.startswith('**') and linea.endswith('**'):
                # Texto en negrita
                p = self.doc.add_paragraph()
                p.style = 'TextoNormal'
                run = p.add_run(linea[2:-2])
                run.bold = True
                i += 1
                continue
            
            # Advertencias o notas especiales
            elif linea.startswith('*⚠️') or linea.startswith('*Nota:'):
                p = self.doc.add_paragraph()
                p.style = 'CitaTextual'
                texto = linea[1:] if linea.startswith('*') else linea
                p.add_run(texto)
                i += 1
                continue
            
            else:
                # Texto normal - VERSIÓN SEGURA
                p = self.doc.add_paragraph()
                p.style = 'TextoNormal'
                self._agregar_texto_con_formato(p, linea)
                i += 1
                continue
        
        if iteraciones >= max_iteraciones:
            print(f"⚠️ ADVERTENCIA: Procesamiento limitado por seguridad en {max_iteraciones} iteraciones")
        
        print(f"🔍 DEBUG - Markdown procesado en {iteraciones} iteraciones")

    def _agregar_texto_con_formato(self, paragraph, texto: str):
        """Agrega texto con formato inline (negritas, cursivas) preservado"""
        # Expresión regular para encontrar texto en negritas **texto**
        patron_negrita = r'\*\*(.*?)\*\*'
        
        # Dividir texto por partes con y sin formato
        partes = re.split(patron_negrita, texto)
        
        for i, parte in enumerate(partes):
            if not parte:
                continue
                
            if i % 2 == 0:
                # Texto normal
                paragraph.add_run(parte)
            else:
                # Texto en negrita
                run = paragraph.add_run(parte)
                run.bold = True

    def _agregar_footer(self):
        """Agrega footer con información de generación"""
        self.doc.add_page_break()
        
        # Footer informativo
        footer_p = self.doc.add_paragraph()
        footer_p.style = 'CitaTextual'
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        fecha_generacion = datetime.now().strftime('%d de %B de %Y')
        footer_text = f"""
Este documento fue generado automáticamente por el Sistema de Análisis Optimia - SCJN.
Fecha de generación: {fecha_generacion}
Revisión y validación jurídica requerida antes de su uso oficial.
        """.strip()
        
        footer_p.add_run(footer_text)

def convertir_markdown_a_word(archivo_markdown: Path, expediente: str, 
                             carpeta_salida: Path, contexto_caso: dict,
                             nombre_ministro: str, nombre_secretario: str, 
                             nombre_secretario_aux: str = None,
                             colaboradores: list = None) -> Path:
    """
    Función utilitaria para convertir un archivo Markdown a Word
    
    Args:
        archivo_markdown: Ruta del archivo .md a convertir
        expediente: Número de expediente
        carpeta_salida: Carpeta donde guardar el archivo Word
        
    Returns:
        Path: Ruta del archivo Word generado
    """
    if not archivo_markdown.exists():
        raise FileNotFoundError(f"Archivo Markdown no encontrado: {archivo_markdown}")
    
    # Leer contenido Markdown
    with open(archivo_markdown, 'r', encoding='utf-8') as f:
        contenido_markdown = f.read()
    
    # Crear nombre de archivo Word
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    nombre_word = f"proyecto_sentencia_{expediente}_{timestamp}.docx"
    archivo_word = carpeta_salida / nombre_word
    
    print(f"🔍 DEBUG Word - Datos recibidos:")
    print(f"  - carpeta_salida: {carpeta_salida}")
    print(f"  - nombre_word: {nombre_word}")
    print(f"  - archivo_word completo: {archivo_word}")
    print(f"  - carpeta_salida existe? {carpeta_salida.exists()}")

    # Generar documento Word
    generator = SCJNWordGenerator()
    print(f"🔍 DEBUG - Guardando Word en: {archivo_word}")
    generator.crear_documento_desde_markdown(
        contenido_markdown=contenido_markdown,
        expediente=expediente,
        archivo_salida=archivo_word,
        contexto_caso=contexto_caso,
        nombre_ministro=nombre_ministro,
        nombre_secretario=nombre_secretario,
        nombre_secretario_aux=nombre_secretario_aux,
        colaboradores=colaboradores
    )
    
    print(f"🔍 DEBUG - Word guardado exitosamente: {archivo_word.exists()}")
    return archivo_word
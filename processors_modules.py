# processors/pdf_processor.py
"""
Procesador para archivos PDF
"""

import base64
from pathlib import Path
from typing import Tuple

class PDFProcessor:
    def extraer_contenido(self, ruta_archivo: Path) -> Tuple[str, str]:
        """
        Extrae contenido de un archivo PDF
        
        Returns:
            Tuple[str, str]: (contenido_base64, tipo_contenido)
        """
        with open(ruta_archivo, 'rb') as f:
            contenido_bytes = f.read()
        
        contenido_base64 = base64.b64encode(contenido_bytes).decode('utf-8')
        return contenido_base64, "pdf"


# processors/docx_processor.py
"""
Procesador para archivos DOCX y DOC
"""

import docx
from pathlib import Path
from typing import Tuple

class DOCXProcessor:
    def extraer_contenido(self, ruta_archivo: Path) -> Tuple[str, str]:
        """
        Extrae contenido de un archivo DOCX/DOC
        
        Returns:
            Tuple[str, str]: (contenido_texto, tipo_contenido)
        """
        try:
            doc = docx.Document(ruta_archivo)
            texto_completo = []
            
            # Extraer párrafos
            for para in doc.paragraphs:
                texto_completo.append(para.text)
            
            # Extraer tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        texto_completo.append(cell.text)
            
            contenido = "\n".join(texto_completo)
            return contenido, "text"
            
        except Exception as e:
            raise Exception(f"Error procesando documento DOCX: {str(e)}")


# processors/txt_processor.py
"""
Procesador para archivos de texto plano
"""

from pathlib import Path
from typing import Tuple
import chardet

class TXTProcessor:
    def extraer_contenido(self, ruta_archivo: Path) -> Tuple[str, str]:
        """
        Extrae contenido de un archivo de texto
        
        Returns:
            Tuple[str, str]: (contenido_texto, tipo_contenido)
        """
        try:
            # Detectar encoding
            with open(ruta_archivo, 'rb') as f:
                raw_data = f.read()
            
            encoding_info = chardet.detect(raw_data)
            encoding = encoding_info.get('encoding', 'utf-8')
            
            # Leer archivo con encoding detectado
            with open(ruta_archivo, 'r', encoding=encoding) as f:
                contenido = f.read()
            
            return contenido, "text"
            
        except Exception as e:
            # Fallback a utf-8
            try:
                with open(ruta_archivo, 'r', encoding='utf-8') as f:
                    contenido = f.read()
                return contenido, "text"
            except:
                raise Exception(f"Error procesando archivo de texto: {str(e)}")


# processors/image_processor.py
"""
Procesador para archivos de imagen (OCR)
"""

import base64
from pathlib import Path
from typing import Tuple

class ImageProcessor:
    def extraer_contenido(self, ruta_archivo: Path) -> Tuple[str, str]:
        """
        Extrae contenido de un archivo de imagen
        
        Returns:
            Tuple[str, str]: (contenido_base64, tipo_contenido)
        """
        try:
            with open(ruta_archivo, 'rb') as f:
                contenido_bytes = f.read()
            
            contenido_base64 = base64.b64encode(contenido_bytes).decode('utf-8')
            return contenido_base64, "image"
            
        except Exception as e:
            raise Exception(f"Error procesando imagen: {str(e)}")


# processors/__init__.py
"""
Módulo de procesadores de documentos
"""

from .pdf_processor import PDFProcessor
from .docx_processor import DOCXProcessor
from .txt_processor import TXTProcessor
from .image_processor import ImageProcessor

__all__ = [
    'PDFProcessor',
    'DOCXProcessor', 
    'TXTProcessor',
    'ImageProcessor'
]